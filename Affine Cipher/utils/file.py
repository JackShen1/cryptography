from docx import Document


def read_docx(path: str) -> str:
    doc = Document(path)
    return "".join([p.text for p in doc.paragraphs])


def write_docx(path: str, prg: str) -> None:
    doc = Document()
    doc.add_paragraph(prg)
    doc.save(path)
