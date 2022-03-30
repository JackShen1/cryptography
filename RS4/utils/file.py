from docx import Document


def read_docx(path: str) -> str:
    doc = Document(path)

    return "".join([p.text for p in doc.paragraphs])


def write_docx(path: str, text: str) -> str:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)

    return text


def read_txt(path: str) -> str:
    with open(path, encoding="utf-8") as in_file:
        return in_file.read()


def write_txt(path: str, text: str) -> str:
    with open(path, "w", encoding="utf-8") as out_file:
        out_file.write(text)

    return text
